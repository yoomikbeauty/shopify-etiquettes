# Import des bibliothèques nécessaires
import textwrap
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import streamlit as st  # pour l'interface web
import requests  # pour faire des requêtes HTTP vers l'API Shopify
import pandas as pd  # pour manipuler les données sous forme de tableaux
import time  # pour ajouter des pauses entre les requêtes
import re  # pour lire la pagination
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from io import BytesIO

# --- QUDO TXT parsing ---------------------------------------------------------
import re
import pandas as pd

def parse_qudo_text_to_df(text: str, include_samples: bool = False) -> pd.DataFrame:
    """
    Attend des lignes du type:
      1 (8809738316993) Beauty of Joseon - Red Bean Water Gel 100ml pcs 15 8.00 120.00 0.00
      2 (MOSTRE) Sample cream 2ml pcs 1 0.01 0.01 0.00
    Retourne: Product Name, Barcode, Unit, Qty, Unit Price EUR, Line Value EUR, VAT EUR
    """
    s = (text or "").replace("–", "-").replace("—", "-")
    s = re.sub(r"[ \t]+", " ", s).strip()

    pat = re.compile(
        r"""
        \b\d+\s*                                  # index
        \((?P<bc>\d{8,14}|MOSTRE)\)\s+            # (barcode|MOSTRE)
        (?P<name>.+?)\s+                          # nom (non-gourmand)
        (?P<unit>pcs|set|box|ea)\s+               # unité
        (?P<qty>\d+)\s+                           # quantité
        (?P<unit_price>\d+[.,]?\d*)\s+            # prix unitaire
        (?P<value>\d+[.,]?\d*)\s+                 # montant ligne
        (?P<vat>\d+[.,]?\d*)                      # tva
        """,
        re.IGNORECASE | re.VERBOSE | re.DOTALL
    )

    rows = []
    for m in pat.finditer(s):
        bc = m.group("bc").strip()
        if (not include_samples) and bc.upper() == "MOSTRE":
            continue

        def fnum(x): return float(str(x).replace(",", "."))
        rows.append({
            "Product Name": m.group("name").strip(" -"),
            "Barcode": bc,
            "Unit": m.group("unit").lower(),
            "Qty": int(m.group("qty")),
            "Unit Price EUR": fnum(m.group("unit_price")),
            "Line Value EUR": fnum(m.group("value")),
            "VAT EUR": fnum(m.group("vat")),
        })
    return pd.DataFrame(rows)

def parse_qudo_name(raw: str, default_vendor: str = "") -> dict:
    """
    Découpe 'Vendor - Title Size' (ou juste 'Vendor Title Size').
    Extrait Vendor, Title, Size (ml/g/oz/kg/l/cl/mg).
    """
    s = str(raw or "").strip().replace("–", "-").replace("—", "-")
    vendor = default_vendor
    # vendor si "Vendor - Title"
    m_dash = re.match(r"^\s*([^-\[\]]{2,}?)\s*-\s*(.+)$", s)
    if m_dash:
        vendor = m_dash.group(1).strip()
        main = m_dash.group(2).strip()
    else:
        main = s

    size_pat = r"(\d+(?:[.,]\d+)?)\s*(ml|g|kg|l|cl|mg|oz)\b"
    sizes = list(re.finditer(size_pat, main, flags=re.I))
    size = ""
    if sizes:
        q, u = sizes[-1].group(1), sizes[-1].group(2)
        size = f"{q.replace(',', '.')} {u.lower()}"
        start, end = sizes[-1].span()
        if end == len(main) or re.match(r"\s*$", main[end:]):
            main = main[:start].strip()

    # title-case doux
    small = {'de','du','des','la','le','les','et','ou','à','au','aux','the','of','for','and','in','on','with'}
    words, out = main.split(), []
    for i,w in enumerate(words):
        if w.isupper() and len(w) <= 4: out.append(w)
        elif w.lower() in small and i not in (0, len(words)-1): out.append(w.lower())
        else: out.append(w.capitalize())
    title = " ".join(out).strip()

    return {"Vendor": vendor, "Title": title, "Size": size}

import math
import re

def price_rounding(raw_price: float, mode: str):
    """Arrondis classiques pour calcul de PV conseillé."""
    if raw_price is None:
        return None
    if mode == ".90 (vers le bas)":
        euros = math.floor(raw_price)
        cents = raw_price - euros
        if cents >= 0.90:
            price = euros + 0.90
        else:
            price = (euros - 1) + 0.90 if euros > 0 else 0.90
        return round(price, 2)
    elif mode == "0,10 le + proche":
        return round(round(raw_price * 10) / 10.0, 2)
    elif mode == ".95 (vers le bas)":
        euros = math.floor(raw_price)
        cents = raw_price - euros
        if cents >= 0.95:
            price = euros + 0.95
        else:
            price = (euros - 1) + 0.95 if euros > 0 else 0.95
        return round(price, 2)
    elif mode == "arrondi sup. à 0,05":
        return round(math.ceil(raw_price * 20) / 20.0, 2)
    else:
        return round(raw_price, 2)

def parse_weight_to_grams(val):
    """
    Parse '208g', '0.2 kg', '7 oz', '1 lb' -> grammes (float).
    """
    if pd.isna(val):
        return None
    s = str(val).strip()
    m = re.search(r"(\d+(?:[.,]\d+)?)\s*(kg|g|oz|lb)\b", s, flags=re.I)
    if not m:
        return None
    q = float(m.group(1).replace(",", "."))
    unit = m.group(2).lower()
    if unit == "kg":
        return q * 1000.0
    if unit == "g":
        return q
    if unit == "lb":
        return q * 453.59237
    if unit == "oz":
        return q * 28.349523125
    return None



from html.parser import HTMLParser

class DocxHTMLParser(HTMLParser):
    def __init__(self, paragraph):
        super().__init__()
        self.paragraph = paragraph
        self.bold = False
        self.italic = False
        self.underline = False

    def handle_starttag(self, tag, attrs):
        if tag == 'b':
            self.bold = True
        elif tag == 'i':
            self.italic = True
        elif tag == 'u':
            self.underline = True

    def handle_endtag(self, tag):
        if tag == 'b':
            self.bold = False
        elif tag == 'i':
            self.italic = False
        elif tag == 'u':
            self.underline = False

    def handle_data(self, data):
        run = self.paragraph.add_run(data)
        run.bold = self.bold
        run.italic = self.italic
        run.underline = self.underline

class PDFTextHTMLParser(HTMLParser):
    def __init__(self, canvas, x, y, font, font_size):
        super().__init__()
        self.c = canvas
        self.x = x
        self.y = y
        self.font = font
        self.font_size = font_size
        self.bold = False
        self.x_cursor = x

    def handle_starttag(self, tag, attrs):
        if tag == 'b':
            self.bold = True

    def handle_endtag(self, tag):
        if tag == 'b':
            self.bold = False

    def handle_data(self, data):
        if self.bold:
            self.c.setFont(self.font + "T", self.font_size)
        else:
            self.c.setFont(self.font, self.font_size)

        self.c.drawString(self.x_cursor, self.y, data)
        text_width = self.c.stringWidth(data, self.font + ("T" if self.bold else ""), self.font_size)
        self.x_cursor += text_width



PRECAUTION_DEFAULT = (
    "<b>Avertissement!</b> Usage externe uniquement. Éviter tout contact avec les yeux. "
    "Tenir hors de portée des enfants. En cas d'apparition de rougeurs, de gonflements ou de démangeaisons pendant ou après l'utilisation, consultez un médecin. "
    "<br><b>A consommer de préférence avant le / Numéro de lot :</b> indiqué sur l'emballage."
)

INFO_BLOCK_TEMPLATE = (
    "<b>Fabricant :</b> {vendor}<br>"
    "<b>EU RP :</b>  Yoomi k-beauty, 19 rue merciere, 68100 Mulhouse, France - 03 65 67 40 62 - SIREN 932 945 256<br>"
    "<b>Fabriqué en Corée</b>"
)

@st.cache_data(ttl=300)
def preparer_stock_csv(csv_path_or_obj, shop_url, access_token):
    df_fournisseur = pd.read_csv(csv_path_or_obj)

    def extraire_barcode(nom):
        match = re.search(r'barcode[\s:-]*([\d]{8,14})', str(nom), re.IGNORECASE)
        return match.group(1) if match else None

    df_fournisseur['Barcode'] = df_fournisseur['Product Name'].apply(extraire_barcode)

    df_fournisseur['Qty'] = (
        df_fournisseur['Qty']
        .astype(str)
        .str.extract(r'(\d+)')
        .fillna(0)
        .astype(int)
    )

    headers = {"X-Shopify-Access-Token": access_token}
    df_variants = get_all_shopify_variants(shop_url, access_token)
    df_merged = pd.merge(df_fournisseur, df_variants, on="Barcode", how="left")

    # 📍 Récupération emplacement (1 seule fois)
    loc_resp = requests.get(f"https://{shop_url}/admin/api/2023-10/locations.json", headers=headers)
    if loc_resp.ok:
        location_id = loc_resp.json()["locations"][0]["id"]
    else:
        location_id = None

    stock_actuels = []
    for i, row in df_merged.iterrows():
        if pd.isna(row["Inventory Item ID"]) or location_id is None:
            stock_actuels.append(None)
            continue
        time.sleep(0.6)  # protection quota
        inv_url = f"https://{shop_url}/admin/api/2023-10/inventory_levels.json"
        params = {"inventory_item_ids": int(row["Inventory Item ID"]), "location_ids": location_id}
        inv_resp = requests.get(inv_url, headers=headers, params=params)
        if inv_resp.ok:
            inv_data = inv_resp.json().get("inventory_levels", [])
            stock_actuels.append(inv_data[0]["available"] if inv_data else 0)
        else:
            stock_actuels.append(None)

    df_merged["Stock actuel"] = stock_actuels
    df_merged["location_id"] = location_id
    return df_merged


# Configuration de la page Streamlit
st.set_page_config(page_title="Shopify Product Viewer", layout="wide")

# Définir la couleur de fond avec du CSS inline


# Enregistrement des polices personnalisées
pdfmetrics.registerFont(TTFont("NotoSans-Italic", "fonts/NotoSans-Italic.ttf"))
pdfmetrics.registerFont(TTFont("AdobeSansMM", "fonts/adobe-sans-mm.ttf"))
pdfmetrics.registerFont(TTFont("IbarraRealNova-Bold", "fonts/IbarraRealNova-Bold.ttf"))
pdfmetrics.registerFont(TTFont("IbarraRealNova-Regular", "fonts/IbarraRealNova-Regular.ttf"))
pdfmetrics.registerFont(TTFont("IbarraRealNova-SemiBold", "fonts/IbarraRealNova-SemiBold.ttf"))
pdfmetrics.registerFont(TTFont("BellCentennial", "fonts/BellCentennialStd-Address.ttf"))
pdfmetrics.registerFont(TTFont("BellCentennialName", "fonts/BellCentennialStd-NameNum.ttf"))



# Titre principal affiché sur la page
st.image("images/logo.png", width=250)
st.markdown("<h1 style='text-align:center'>Créateur de carte YOOMI</h1>", unsafe_allow_html=True)


tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
    "Base de données", "Étiquettes prix", "Étiquettes de traduction Fournisseur",
    "Étiquettes de traduction Boutique", "📦 Stock fournisseur",
    "📦 Gestion stock manuels", "💸 Gestion Soldes",
    "➕ Nouveaux produits (CSV fournisseur)"
])


with tab1:
    st.markdown("## Base de données produits")
    # Bouton pour mettre à jour les données Shopify
    with st.expander("🛠 Paramètres de récupération de la base de données"):
        import os
        # 👉 Shopify credentials via Streamlit Cloud secrets
        shop_url = st.secrets["shopify"]["shop_url"]
        access_token = st.secrets["shopify"]["access_token"]
        mode_complet = st.checkbox("Inclure les métadonnées personnalisées (plus lent)", value=True)
        only_recent = st.checkbox("Afficher uniquement les 50 derniers produits ajoutés")
        force_update = st.checkbox("🔁 Forcer une mise à jour complète (ignorer les dates)", value=False)
        st.session_state.force_update = force_update
        st.success(f"🔐 Connecté à {shop_url}")

    # Chargement initial depuis fichier CSV si existant
    if os.path.exists("data/produits_shopify.csv"):
        try:
            df_temp = pd.read_csv("data/produits_shopify.csv")
            if "updated_at" in df_temp.columns:
                last_update_display = pd.to_datetime(df_temp["updated_at"], errors="coerce").max()
                total_count = len(df_temp)
                st.markdown(f"<div style='text-align:center; margin-top:10px; font-size:14px; color:#333;'>❤️ Dernière mise à jour : <b>{last_update_display.strftime('%d/%m/%Y %H:%M')}</b> — {total_count} produits enregistrés</div>", unsafe_allow_html=True)
        except:
            pass
    if "df" not in st.session_state:
        if os.path.exists("data/produits_shopify.csv"):
            st.session_state['df'] = pd.read_csv("data/produits_shopify.csv")
            st.success("Base produits chargée depuis le fichier local.")

    # Mise à jour manuelle
    last_updated = None
    if os.path.exists("data/produits_shopify.csv"):
            try:
                old_df = pd.read_csv("data/produits_shopify.csv")
                if "updated_at" in old_df.columns:
                    last_updated = pd.to_datetime(old_df["updated_at"], errors="coerce").max()
            except:
                last_updated = None
    if st.button("Mettre à jour la base produits depuis Shopify"):
        st.info("Connexion à Shopify...")

        headers = {
            "X-Shopify-Access-Token": access_token
        }

        base_url = f"https://{shop_url}/admin/api/2023-10/products.json"
        metafield_url_template = f"https://{shop_url}/admin/api/2023-10/products/{{product_id}}/metafields.json"

        products = []
        page_info = None

        with st.spinner("Chargement des produits..."):
            params = {"limit": 250, "order": "updated_at asc"}
            if last_updated is not None and not st.session_state.get('force_update', False):
                params["updated_at_min"] = last_updated.isoformat() 

            while True:
                response = requests.get(base_url, headers=headers, params=params)
                response.raise_for_status()
                batch = response.json().get("products", [])
                products.extend([p for p in batch if p.get("status") == "active"])

                if only_recent:
                    break

                link_header = response.headers.get("Link")
                if link_header and 'rel="next"' in link_header:
                    match = re.search(r'<[^>]*[?&]page_info=([^&>]*)[^>]*>; rel="next"', link_header)
                    if match:
                        page_info = match.group(1)
                        params = {"limit": 250, "page_info": page_info}
                    else:
                        break
                else:
                    break

        if not products:
            st.warning("Aucun produit trouvé.")
        else:
            data = []
            all_new_data = []
            if mode_complet:
                metafield_keys = [
                    "mini_description", "moyenne_description", "utilisation", "taille","ingredients", "routine",
                    "info_bestseller", "info_cruelty_free", "info_vegan", "info_clean_beauty",
                    "tout_type", "peau_grasse", "peau_mature", "peau_seche", "peau_sensible", "peau_acneique", "periode_mois", "texte_recyclage"
                    
                ]

                progress_bar = st.progress(0)
                status_text = st.empty()

                for i, p in enumerate(products):
                    product_id = p.get("id")
                    title = p.get("title")
                    status_text.text(f"Récupération des métadonnées pour : {title} (ID {product_id})")
                    meta_fail = False

                    time.sleep(0.8)
                    metafields_response = requests.get(metafield_url_template.format(product_id=product_id), headers=headers)
                    metafields = metafields_response.json().get("metafields", [])
                    metafield_data = {key: "" for key in metafield_keys}

                    for meta in metafields:
                        key = meta.get("key")
                        if meta.get("namespace") == "custom" and key in metafield_keys:
                            value = meta.get("value")
                            # Vérification + Retry si value vide
                            retry_count = 0
                            while (value is None or value == "") and retry_count < 3:
                                time.sleep(0.7)
                                retry_response = requests.get(metafield_url_template.format(product_id=product_id), headers=headers)
                                retry_meta = retry_response.json().get("metafields", [])
                                for retry_item in retry_meta:
                                    if retry_item.get("namespace") == "custom" and retry_item.get("key") == key:
                                        value = retry_item.get("value")
                                        break
                                retry_count += 1
                            if value:
                                metafield_data[key] = str(value)
                            else:
                                meta_fail = True

                    all_new_data.append({
                        "ID": p.get("id"),
                        "updated_at": p.get("updated_at"),
                        "Vendor": p.get("vendor"),
                        "Title": title,
                        "Type": p.get("product_type"),
                        "Variant Price": p.get("variants", [{}])[0].get("price"),
                        "Variant Compare Price": p.get("variants", [{}])[0].get("compare_at_price"),
                        "Variant Barcode": p.get("variants", [{}])[0].get("barcode"),
                        **{f"custom.{key}": metafield_data[key] for key in metafield_keys}
                    })
                    progress_bar.progress((i + 1) / len(products))
                status_text.text("Récupération terminée.")
                if meta_fail:
                    st.warning(f"⚠️ Certains produits n'ont pas toutes leurs métadonnées. Veuillez vérifier manuellement.")

            else:
                for p in products:
                    data.append({
                        "Vendor": p.get("vendor"),
                        "Title": p.get("title"),
                        "Type": p.get("product_type"),
                        "Variant Price": p.get("variants", [{}])[0].get("price"),
                        "Variant Barcode": p.get("variants", [{}])[0].get("barcode")
                    })

            data = sorted(data, key=lambda x: x["ID"], reverse=True)
            df = pd.DataFrame(all_new_data)
            if os.path.exists("data/produits_shopify.csv"):
                try:
                    old_df = pd.read_csv("data/produits_shopify.csv")
                    combined_df = pd.concat([old_df, df], ignore_index=True)
                    df = combined_df.drop_duplicates(subset="ID", keep="last")
                except:
                    pass
            df.to_csv("data/produits_shopify.csv", index=False)
            st.session_state['df'] = df
            st.success(f"{len(df)} produits récupérés et enregistrés dans 'data/produits_shopify.csv'.")

    # Affichage + export CSV + sélection PDF si données présentes
    if 'df' in st.session_state:
        df = st.session_state['df']
        with st.expander("### 🔍 Filtrer à partir d’un fichier de commande"):
            source_cmd = st.selectbox(
                "Source du bon",
                ["StyleKorean (CSV)", "QUDO (TXT)"],
                key="src_cmd_tab1"
            )

            df_filtered_by_cmd = None
            barcodes_non_trouves = []

            if source_cmd == "StyleKorean (CSV)":
                commande_csv = st.file_uploader("📁 Uploader le fichier CSV StyleKorean", type=["csv"], key="cmd_csv_tab1")
                if commande_csv:
                    try:
                        df_commande = pd.read_csv(commande_csv)

                        def extraire_barcode(nom):
                            m = re.search(r'barcode[\s:-]*([\d]{8,14})', str(nom), re.IGNORECASE)
                            return m.group(1) if m else None

                        df_commande["extracted_barcode"] = df_commande["Product Name"].apply(extraire_barcode)
                        barcodes_commande = df_commande["extracted_barcode"].dropna().unique().tolist()

                        if barcodes_commande:
                            df_barcodes = df["Variant Barcode"].astype(str)
                            barcodes_trouves = df_barcodes[df_barcodes.isin(barcodes_commande)].unique().tolist()
                            barcodes_non_trouves = sorted(set(barcodes_commande) - set(barcodes_trouves))
                            df = df[df["Variant Barcode"].astype(str).isin(barcodes_commande)]
                            st.success(f"{len(df)} produits trouvés (sur {len(barcodes_commande)} barcodes du bon).")
                        else:
                            st.warning("Aucun code-barres valide trouvé dans le CSV.")
                    except Exception as e:
                        st.error(f"Erreur lecture CSV : {e}")

            else:  # QUDO (TXT)
                commande_txt = st.file_uploader("📄 Uploader le fichier texte QUDO", type=["txt"], key="cmd_txt_tab1")
                if commande_txt:
                    try:
                        content = commande_txt.read().decode("utf-8", errors="ignore")
                        df_txt = parse_qudo_text_to_df(content, include_samples=False)
                        st.dataframe(df_txt[["Product Name","Barcode","Qty"]], use_container_width=True)
                        barcodes_commande = df_txt["Barcode"].astype(str).unique().tolist()

                        df_barcodes = df["Variant Barcode"].astype(str)
                        barcodes_trouves = df_barcodes[df_barcodes.isin(barcodes_commande)].unique().tolist()
                        barcodes_non_trouves = sorted(set(barcodes_commande) - set(barcodes_trouves))
                        df = df[df["Variant Barcode"].astype(str).isin(barcodes_commande)]
                        st.success(f"{len(df)} produits trouvés (sur {len(barcodes_commande)} barcodes du bon).")
                    except Exception as e:
                        st.error(f"Erreur lecture TXT : {e}")

            if barcodes_non_trouves:
                with st.expander("⚠️ Barcodes non trouvés dans Shopify"):
                    for bc in barcodes_non_trouves:
                        st.markdown(f"- `{bc}`")

        st.dataframe(df, use_container_width=True)
        # 👉 Sauver ce qui est VRAIMENT affiché en tab1 pour réutilisation ailleurs
        st.session_state["df_view_tab1"] = df.copy()

        csv = df.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="Télécharger en CSV",
            data=csv,
            file_name="data/produits_shopify.csv",
            mime="text/csv",
        )



        with tab2:
            st.markdown("## Création d’étiquettes prix")
            # Helpers anti-"nan"
            def filled(v):
                return pd.notna(v) and str(v).strip() != '' and str(v).strip().lower() != 'nan'

            def text(v):
                return str(v).strip() if filled(v) else ''

            # Choix des produits à étiqueter
            st.markdown("### Étiquettes à imprimer")
            # Label lisible même si certaines colonnes sont vides
            tailles = df['custom.taille'].apply(lambda x: f" ({text(x)})" if filled(x) else '')
            vendors = df['Vendor'].fillna('').astype(str)
            titles  = df['Title'].fillna('').astype(str)

            df['label'] = (vendors + ' - ' + titles + tailles).str.replace(r'^\s*-\s*', '', regex=True).str.strip()
            df = df.sort_values('ID', ascending=False).reset_index(drop=True)

            selected_labels = st.multiselect(
                "Sélectionnez les produits à imprimer : (8 max par page)",
                options=df['label'].tolist(),
                placeholder="Choisissez un ou plusieurs produits..."
            )
            filtered_df = df[df['label'].isin(selected_labels)].reset_index(drop=True)

            if st.button("Générer les étiquettes PDF (8 par page)") and not filtered_df.empty:
                buffer = BytesIO()
                c = canvas.Canvas(buffer, pagesize=A4)

                width, height = A4
                margin_x, margin_y = 19 * mm, 38.52 * mm
                label_width, label_height = 86 * mm, 55 * mm

                for i, row in filtered_df.iterrows():
                    if i > 0 and i % 8 == 0:
                        c.showPage()

                    col = i % 2
                    row_pos = (i // 2) % 4
                    x = margin_x + col * label_width
                    y = height - margin_y - (row_pos + 1) * label_height

                    # CADRE
                    c.setFont("Helvetica", 8)
                    c.setStrokeColorRGB(0, 0, 0)
                    c.rect(x, y, label_width, label_height, stroke=0, fill=0)

                    # VENDOR
                    c.setFont("IbarraRealNova-SemiBold", 11.5)
                    if filled(row.get('Vendor')):
                        vendor_text = c.beginText()
                        vendor_text.setTextOrigin(x + 2.5 * mm, y + label_height - 5 * mm)
                        for line in textwrap.wrap(text(row.get('Vendor')), width=28)[:2]:
                            vendor_text.textLine(line)
                        c.drawText(vendor_text)

                    # Séparateur après Vendor
                    c.setLineWidth(0.1)
                    c.line(x, y + label_height - 6.5 * mm, x + label_width, y + label_height - 6.5 * mm)

                    # TAILLE (à droite, si dispo)
                    if filled(row.get('custom.taille')):
                        c.setFont("IbarraRealNova-SemiBold", 11.5)
                        c.drawRightString(x + label_width - 2.5 * mm, y + label_height - 5 * mm, text(row.get('custom.taille')))

                    # TITRE
                    if filled(row.get('Title')):
                        c.setFont("IbarraRealNova-Bold", 14)
                        wrapped_title = textwrap.wrap(text(row.get('Title')), width=30)
                        line_height = 13
                        total_lines = len(wrapped_title[:2])
                        for idx, line in enumerate(wrapped_title[:2]):
                            text_width = pdfmetrics.stringWidth(line, "IbarraRealNova-Bold", 14)
                            y_offset = y + label_height - 14 * mm if total_lines == 1 else y + label_height - 12 * mm - (idx * line_height)
                            c.drawString(x + (label_width - text_width) / 2, y_offset, line)

                    # Séparateur après Titre
                    c.setLineWidth(0.1)
                    c.line(x, y + label_height - 18 * mm, x + label_width, y + label_height - 18 * mm)

                    # DESCRIPTION
                    if filled(row.get('custom.moyenne_description')):
                        c.setFont("AdobeSansMM", 7.5)
                        desc = text(row.get('custom.moyenne_description'))
                        line_height = 10.2
                        for idx, line in enumerate(textwrap.wrap(desc, width=61.5)):
                            c.drawString(x + 2.5 * mm, y + label_height - 22 * mm - (idx * line_height), line)

                    # Séparateur (gris)
                    c.setLineWidth(0.1)
                    c.setStrokeColorRGB(0.7, 0.7, 0.7)
                    c.line(x, y + label_height - 38 * mm, x + label_width, y + label_height - 38 * mm)
                    c.setStrokeColorRGB(0, 0, 0)

                    # ROUTINE
                    if filled(row.get('custom.routine')):
                        c.setFont("IbarraRealNova-Regular", 7)
                        c.setFillColorRGB(0.4, 0.4, 0.4)
                        c.drawRightString(x + label_width - 2.5 * mm, y + 9 * mm, f"Étape n° {text(row.get('custom.routine'))}")
                        c.setFillColorRGB(0, 0, 0)

                    # ICONES CONDITIONNELLES
                    icon_size = 12 * mm
                    icon_y = y + label_height - 53 * mm
                    icon_x = x + 2 * mm
                    if str(row.get('custom.info_vegan', '')).strip().lower() == 'true':
                        try:
                            c.drawImage("images/vegan.png", icon_x, icon_y, width=icon_size, height=icon_size, preserveAspectRatio=True)
                            icon_x += icon_size
                        except Exception:
                            c.setFillColorRGB(1, 0, 0); c.rect(icon_x, icon_y, icon_size, icon_size, fill=1); icon_x += icon_size
                    if str(row.get('custom.info_cruelty_free', '')).strip().lower() == 'true':
                        try:
                            c.drawImage("images/cruelty.png", icon_x, icon_y, width=icon_size, height=icon_size, preserveAspectRatio=True)
                            icon_x += icon_size
                        except Exception:
                            c.setFillColorRGB(1, 0, 0); c.rect(icon_x, icon_y, icon_size, icon_size, fill=1); icon_x += icon_size
                    if str(row.get('custom.info_clean_beauty', '')).strip().lower() == 'true':
                        try:
                            c.drawImage("images/clean.png", icon_x, icon_y, width=icon_size, height=icon_size, preserveAspectRatio=True)
                            icon_x += icon_size
                        except Exception:
                            c.setFillColorRGB(1, 0, 0); c.rect(icon_x, icon_y, icon_size, icon_size, fill=1); icon_x += icon_size

                    # TYPE DE PEAU / CHEVEUX
                    types = []
                    t = text(row.get('Type'))
                    if t == 'P':
                        if str(row.get('custom.tout_type', '')).strip().lower() in ['true', '1']: types.append("• tout type")
                        if str(row.get('custom.peau_acneique', '')).strip().lower() in ['true', '1']: types.append("• acnéique")
                        if str(row.get('custom.peau_grasse', '')).strip().lower() in ['true', '1']: types.append("• grasse")
                        if str(row.get('custom.peau_seche', '')).strip().lower() in ['true', '1']: types.append("• sèche")
                        if str(row.get('custom.peau_sensible', '')).strip().lower() in ['true', '1']: types.append("• sensible")
                        if str(row.get('custom.peau_mature', '')).strip().lower() in ['true', '1']: types.append("• mature")
                        label_type = "Type de peau :"
                    elif t == 'C':
                        if str(row.get('custom.tout_type', '')).strip().lower() in ['true', '1']: types.append("• tout type")
                        if str(row.get('custom.peau_grasse', '')).strip().lower() in ['true', '1']: types.append("• gras")
                        if str(row.get('custom.peau_seche', '')).strip().lower() in ['true', '1']: types.append("• sec")
                        if str(row.get('custom.peau_sensible', '')).strip().lower() in ['true', '1']: types.append("• sensible")
                        label_type = "Type de cheveux :"
                    else:
                        label_type = ""

                    if types:
                        c.setFont("NotoSans-Italic", 9)
                        c.drawRightString(x + label_width - 2.5 * mm, y + label_height - 42 * mm, f"{label_type} {' '.join(types)}")

                    c.setFillColorRGB(0, 0, 0)

                    # PRIX + PRIX BARRÉ
                    if filled(row.get('Variant Price')):
                        try:
                            prix = float(str(row.get('Variant Price')).replace(',', '.'))
                            prix_str = f"{prix:.2f}".replace('.', ',') + "€"

                            affiche_compare = False
                            cp = row.get('Variant Compare Price')
                            if filled(cp):
                                try:
                                    compare_price_float = float(str(cp).replace(',', '.'))
                                    if compare_price_float > prix:
                                        affiche_compare = True
                                        c.setFont("IbarraRealNova-Regular", 14)
                                        compare_price_str = f"{compare_price_float:.2f}".replace('.', ',') + "€"
                                        text_width = pdfmetrics.stringWidth(compare_price_str, "IbarraRealNova-Regular", 10)
                                        compare_price_x = x + label_width - 30 * mm - text_width
                                        compare_price_y = y + 3 * mm
                                        c.setFillColorRGB(0, 0, 0)
                                        c.drawString(compare_price_x, compare_price_y, compare_price_str)
                                        c.setLineWidth(0.5)
                                        c.line(compare_price_x, compare_price_y + 4, compare_price_x + text_width, compare_price_y + 4)
                                        c.setFillColorRGB(0, 0, 0)
                                except Exception:
                                    pass  # mauvaise donnée compare price: on ignore

                            c.setFont("IbarraRealNova-Bold", 20)
                            c.setFont("IbarraRealNova-Bold", 20)
                            if affiche_compare:
                                c.setFillColorRGB(1, 0, 0)
                            else:
                                c.setFillColorRGB(0, 0, 0)
                            c.drawRightString(x + label_width - 2.5 * mm, y + 3 * mm, prix_str)
                            c.setFillColorRGB(0, 0, 0)
                            c.drawRightString(x + label_width - 2.5 * mm, y + 3 * mm, prix_str)
                            c.setFillColorRGB(0, 0, 0)
                        except Exception as e:
                            c.setFont("Helvetica", 8)
                            c.drawString(x + 2 * mm, y + 3 * mm, f"Erreur prix: {e}")

                c.save()
                buffer.seek(0)
                st.download_button(
                    label="Télécharger les étiquettes en PDF",
                    data=buffer.getvalue(),
                    file_name="etiquettes_shopify.pdf",
                    mime="application/pdf",
                )



# === 📦 MISE À JOUR STOCK FOURNISSEUR =======================
with tab5:
    import time
    import requests
    st.markdown("## Mise à jour du stock via bon de commande fournisseur (STYLE KOREAN)")
    csv_fournisseur = st.file_uploader("📁 Uploader le fichier CSV fournisseur", type=["csv"])
    
    # Fonction de prétraitement + cache
    @st.cache_data(ttl=600)
    def preparer_stock_csv(csv_file, shop_url, access_token):
        df_fournisseur = pd.read_csv(csv_file)

        def extraire_barcode(nom):
            match = re.search(r'barcode[\s:-]*([\d]{8,14})', str(nom), re.IGNORECASE)
            return match.group(1) if match else None

        df_fournisseur['Barcode'] = df_fournisseur['Product Name'].apply(extraire_barcode)

        df_fournisseur['Qty'] = (
            df_fournisseur['Qty']
            .astype(str)
            .str.extract(r'(\d+)')
            .fillna(0)
            .astype(int)
        )

        headers = {"X-Shopify-Access-Token": access_token}

        # Récupération des variantes Shopify
        def get_all_variants():
            all_variants = []
            url = f"https://{shop_url}/admin/api/2023-10/products.json?limit=250"
            while url:
                resp = requests.get(url, headers=headers)
                resp.raise_for_status()
                products = resp.json().get("products", [])
                for p in products:
                    for v in p.get("variants", []):
                        all_variants.append({
                            "Product Title": p["title"],
                            "Variant Title": v["title"],
                            "Barcode": v.get("barcode"),
                            "Variant ID": v["id"],
                            "Inventory Item ID": v["inventory_item_id"]
                        })
                link = resp.headers.get("Link", "")
                if 'rel="next"' in link:
                    match = re.search(r'<([^>]+)>; rel="next"', link)
                    url = match.group(1) if match else None
                else:
                    url = None
            return pd.DataFrame(all_variants)

        df_variants = get_all_variants()
        df_merged = pd.merge(df_fournisseur, df_variants, on="Barcode", how="left")

        # Récupération location
        loc_resp = requests.get(f"https://{shop_url}/admin/api/2023-10/locations.json", headers=headers)
        if loc_resp.ok:
            location_id = loc_resp.json()["locations"][0]["id"]
        else:
            location_id = None

        stock_actuels = []
        for i, row in df_merged.iterrows():
            if pd.isna(row["Inventory Item ID"]) or location_id is None:
                stock_actuels.append(None)
                continue
            time.sleep(0.6)  # anti quota
            inv_url = f"https://{shop_url}/admin/api/2023-10/inventory_levels.json"
            params = {"inventory_item_ids": int(row["Inventory Item ID"]), "location_ids": location_id}
            inv_resp = requests.get(inv_url, headers=headers, params=params)
            if inv_resp.ok:
                inv_data = inv_resp.json().get("inventory_levels", [])
                stock_actuels.append(inv_data[0]["available"] if inv_data else 0)
            else:
                stock_actuels.append(None)

        df_merged["Stock actuel"] = stock_actuels
        df_merged["location_id"] = location_id
        return df_merged

    if csv_fournisseur:
        df_merged = preparer_stock_csv(csv_fournisseur, shop_url, access_token)
        st.session_state['df_stock_update'] = df_merged

        st.markdown("### 📊 Aperçu")
        st.dataframe(df_merged[["Product Name", "Barcode", "Stock actuel", "Qty"]], use_container_width=True)

        if st.button("✅ Mettre à jour tous les stocks", key="maj_global"):
            progress_bar = st.progress(0)
            total = len(df_merged)

            for i, row in df_merged.iterrows():
                if pd.isna(row["Inventory Item ID"]) or pd.isna(row["location_id"]):
                    st.warning(f"⚠️ Produit introuvable : {row['Product Name']}")
                    continue

                payload = {
                    "location_id": int(row["location_id"]),
                    "inventory_item_id": int(row["Inventory Item ID"]),
                    "available_adjustment": int(row["Qty"])
                }

                for attempt in range(2):  # retry une fois si trop de requêtes
                    resp = requests.post(
                        f"https://{shop_url}/admin/api/2023-10/inventory_levels/adjust.json",
                        headers={"X-Shopify-Access-Token": access_token},
                        json=payload
                    )

                    if resp.status_code == 200:
                        st.success(f"✔️ {row['Product Name']} → +{row['Qty']}")
                        break
                    elif resp.status_code == 429:
                        st.warning(f"⏳ Trop de requêtes pour : {row['Product Name']} — nouvelle tentative dans 5s...")
                        time.sleep(5)
                    else:
                        st.error(f"❌ Échec : {row['Product Name']} → {resp.status_code}")
                        break

                progress_bar.progress((i + 1) / total)
                time.sleep(0.3)  # délai anti-quota

        # 🔘 MAJ individuelle sans recalcul
        st.markdown("### 🛠 Mise à jour individuelle")
        for i, row in st.session_state.get('df_stock_update', pd.DataFrame()).iterrows():
            with st.expander(f"🔹 {row['Product Name']} — Barcode: {row['Barcode']}"):
                st.write(f"Stock actuel : **{row['Stock actuel']}**")
                st.write(f"Ajouter : **{row['Qty']}**")

                if st.button(f"Mettre à jour ce produit", key=f"btn_indiv_{i}"):
                    payload = {
                        "location_id": int(row["location_id"]),
                        "inventory_item_id": int(row["Inventory Item ID"]),
                        "available_adjustment": int(row["Qty"])
                    }
                    resp = requests.post(
                        f"https://{shop_url}/admin/api/2023-10/inventory_levels/adjust.json",
                        headers={"X-Shopify-Access-Token": access_token},
                        json=payload
                    )
                    if resp.ok:
                        st.success(f"✔️ Stock mis à jour : {row['Product Name']}")
                    else:
                        st.error(f"❌ Erreur : {row['Product Name']}")







with tab6:
    st.markdown("## Mise à jour manuelle du stock par barcode")

    barcode_input = st.text_input("🔍 Entrez le barcode du produit à mettre à jour")
    qty_input = st.number_input("📦 Quantité à ajouter (positive ou négative)", value=0, step=1)

    if st.button("Mettre à jour le stock Shopify"):
        if not barcode_input or qty_input == 0:
            st.warning("Saisis un barcode valide et une quantité différente de 0.")
        else:
            with st.spinner("Connexion à Shopify..."):
                try:
                    # 🔍 Étape 1 : Recherche du produit par barcode
                    search_url = f"https://{shop_url}/admin/api/2023-10/products.json"
                    params = {"fields": "id,title,variants", "limit": 250}
                    headers = {"X-Shopify-Access-Token": access_token}
                    found_variant = None

                    while True:
                        resp = requests.get(search_url, headers=headers, params=params)
                        resp.raise_for_status()
                        products = resp.json()["products"]

                        for p in products:
                            for v in p.get("variants", []):
                                if str(v.get("barcode", "")).strip() == barcode_input.strip():
                                    found_variant = v
                                    break
                            if found_variant:
                                break

                        link = resp.headers.get("Link", "")
                        if 'rel="next"' in link:
                            match = re.search(r'<([^>]+)>; rel="next"', link)
                            if match:
                                search_url = match.group(1)
                            else:
                                break
                        else:
                            break

                    if not found_variant:
                        st.error("❌ Aucun produit trouvé avec ce barcode.")
                    else:
                        variant_id = found_variant["id"]
                        inventory_item_id = found_variant["inventory_item_id"]

                        # 🔧 Étape 2 : Récupérer location_id
                        loc_url = f"https://{shop_url}/admin/api/2023-10/locations.json"
                        loc_resp = requests.get(loc_url, headers=headers)
                        loc_resp.raise_for_status()
                        location_id = loc_resp.json()["locations"][0]["id"]

                        # 🔎 Étape 3 : Afficher stock actuel
                        inv_url = f"https://{shop_url}/admin/api/2023-10/inventory_levels.json"
                        inv_params = {"inventory_item_ids": inventory_item_id, "location_ids": location_id}
                        inv_resp = requests.get(inv_url, headers=headers, params=inv_params)
                        inv_resp.raise_for_status()
                        inv_data = inv_resp.json().get("inventory_levels", [])
                        stock_actuel = inv_data[0]["available"] if inv_data else 0

                        st.info(f"Stock actuel : {stock_actuel} → après mise à jour : {stock_actuel + qty_input}")

                        # 🔁 Étape 4 : Envoyer la mise à jour
                        payload = {
                            "location_id": location_id,
                            "inventory_item_id": inventory_item_id,
                            "available_adjustment": qty_input
                        }
                        update_url = f"https://{shop_url}/admin/api/2023-10/inventory_levels/adjust.json"
                        update_resp = requests.post(update_url, headers=headers, json=payload)
                        update_resp.raise_for_status()

                        st.success("✅ Stock mis à jour avec succès !")
                        st.json(update_resp.json())

                except Exception as e:
                    st.error(f"❌ Erreur : {e}")

with tab3:
    st.markdown("## 📄 Générateur d’étiquettes Word pour la traduction/fournisseur")

    # --- 2 sources possibles ---
    source = st.radio(
        "Source des données",
        ["Depuis un CSV", "Depuis l’onglet 1 (liste filtrée)"],
        horizontal=True
    )

    # --- Petite fonction commune pour fabriquer le DOCX depuis un DataFrame ---
    def build_doc_from_df(df_src: pd.DataFrame) -> BytesIO:
        doc = Document()
        doc.add_paragraph()

        def feed_html(p, html_text):
            parser = DocxHTMLParser(p)
            parser.feed(str(html_text).replace("<br>", "\n"))

        for _, row in df_src.iterrows():
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)

            # Titre
            feed_html(cell.add_paragraph(), f"<b>{row.get('Vendor', '')}</b>\n<b>{row.get('Title', '')}</b>")

            # Contenance
            cont = row.get('custom.taille', '')
            if pd.notna(cont) and str(cont).strip():
                feed_html(cell.add_paragraph(), f"<b>Contenance :</b> {str(cont).strip()}")

            # Barcode
            barcode = str(row.get('Variant Barcode', ''))
            feed_html(cell.add_paragraph(), f"<b>Barcode :</b> {barcode}")

            # Utilisation
            util = str(row.get('custom.utilisation', ''))
            if util and util.lower() != 'nan':
                feed_html(cell.add_paragraph(), f"<b>Mode d'emploi :</b> {util}")

            # Ingrédients
            ing = str(row.get('custom.ingredients', ''))
            feed_html(cell.add_paragraph(), f"<b>Ingrédients :</b> {ing}")

            # Précaution
            feed_html(cell.add_paragraph(), PRECAUTION_DEFAULT)

            # Infos fabricant
            vendor = str(row.get('Vendor', ''))
            feed_html(cell.add_paragraph(), INFO_BLOCK_TEMPLATE.format(vendor=vendor))

            # Icônes
            pao_value = row.get('custom.periode_mois', '')
            if pd.isna(pao_value) or str(pao_value).strip() == '':
                pao_icon = "pao_12m.png"
            else:
                try:
                    pao_int = int(float(pao_value))
                    pao_icon = f"pao_{pao_int}m.png"
                except:
                    pao_icon = "pao_12m.png"

            tri_value = str(row.get('custom.texte_recyclage', '')).strip().lower().replace(' ', '_')
            tri_icon = f"{tri_value}.png" if tri_value not in ['', 'nan'] else "tri_standard.png"

            p = cell.add_paragraph()
            run = p.add_run()

            # (Assure-toi d'avoir import os au top du fichier)
            if os.path.exists(f"icones/{pao_icon}"):
                run.add_picture(f"icones/{pao_icon}", width=Inches(0.6))
            if os.path.exists(f"icones/{tri_icon}"):
                run.add_picture(f"icones/{tri_icon}", width=Inches(2))

            # Bordures
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s>'
                r'<w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                r'<w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                r'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                r'<w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                r'</w:tcBorders>' % nsdecls('w')))

            doc.add_paragraph()

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # --- Branche CSV (inchangé, mais réutilise la fonction commune) ---
    if source == "Depuis un CSV":
        uploaded_csv = st.file_uploader("📁 Fichier produits (CSV)", type=["csv"])
        if uploaded_csv:
            df_csv = pd.read_csv(uploaded_csv)
            buffer = build_doc_from_df(df_csv)
            st.download_button(
                label="📥 Télécharger l'étiquette Word",
                data=buffer.getvalue(),
                file_name="Etiquettes_Produits_YOOMI.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # --- NOUVELLE branche : depuis la liste filtrée de l’onglet 1 ---
    else:
        if "df_view_tab1" not in st.session_state or st.session_state["df_view_tab1"].empty:
            st.warning("⚠️ Aucune liste filtrée détectée. Va d’abord dans l’onglet 1, applique tes filtres, puis reviens ici.")
        else:
            df_src = st.session_state["df_view_tab1"].copy()

            # (optionnel) permettre de restreindre encore via un multiselect local
            df_src['__label__'] = df_src['Vendor'].astype(str) + " - " + df_src['Title'].astype(str)
            subset = st.multiselect(
                "Sélectionne (facultatif) des produits parmi la liste filtrée de l’onglet 1 :",
                options=df_src['__label__'].tolist()
            )
            if subset:
                df_src = df_src[df_src['__label__'].isin(subset)]

            if df_src.empty:
                st.info("La sélection est vide.")
            else:
                buffer = build_doc_from_df(df_src)
                st.download_button(
                    label=f"📥 Télécharger {len(df_src)} étiquette(s) en Word",
                    data=buffer.getvalue(),
                    file_name="Etiquettes_Produits_YOOMI.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen.canvas import Canvas

import base64
from reportlab.lib.pagesizes import portrait
from reportlab.pdfgen.canvas import Canvas
from PIL import Image
import fitz  # PyMuPDF pour prévisualisation PDF

with tab4:
    st.markdown("## 📄 Étiquettes de traduction (5×5 cm) avec polices personnalisées")

    if "df" not in st.session_state:
        st.warning("⚠️ Charge d'abord les produits depuis l’onglet 1.")
    else:
        df = st.session_state["df"]
        df['label'] = df['Vendor'] + ' - ' + df['Title']
        selected_labels = st.multiselect("📌 Sélectionne les produits", df['label'].tolist())

        df_filtered = df[df['label'].isin(selected_labels)].reset_index(drop=True)

        if not df_filtered.empty:
            from reportlab.lib.pagesizes import portrait
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Frame, KeepInFrame, Spacer
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib.units import mm
            from reportlab.lib.colors import black
            from PIL import Image
            import fitz
            import os
            from io import BytesIO
            import pandas as pd
            from reportlab.pdfgen.canvas import Canvas
            from reportlab.platypus.flowables import HRFlowable

            pdfmetrics.registerFont(TTFont("BellCentennial", "fonts/BellCentennialStd-Address.ttf"))
            pdfmetrics.registerFont(TTFont("BellCentennial-Bold", "fonts/BellCentennialStd-NameNum.ttf"))

            title_style = ParagraphStyle('title_style', fontName='Helvetica', fontSize=6, alignment=TA_CENTER, leading=5)
            subtitle_style = ParagraphStyle('subtitle_style', fontName='Helvetica', fontSize=5, alignment=TA_CENTER, leading=5.0)
            text_style = ParagraphStyle('text_style', fontName='Helvetica', fontSize=5, alignment=0, leading=4.8)
            small_text_style = ParagraphStyle('small_text', fontName='Helvetica', fontSize=4.5, alignment=0, leading=4.6)

            for i, row in df_filtered.iterrows():
                buffer = BytesIO()
                pdf = SimpleDocTemplate(buffer, pagesize=(141.73, 141.73), leftMargin=0, rightMargin=0, topMargin=0, bottomMargin=0)

                # Préparer les blocs (Paragraphs)
                title_story = [Paragraph(f"<b>{row.get('Vendor', '')} - {row.get('Title', '')}</b>", title_style)]

                mini_desc = str(row.get("custom.mini_description", ""))
                taille = str(row.get("custom.taille", ""))
                description = f"{mini_desc} - {taille}" if taille and taille.lower() != "nan" else mini_desc
                desc_story = [Paragraph(description, subtitle_style)]

                util = str(row.get("custom.utilisation", ""))
                if util and util.lower() != 'nan':
                    util_para = Paragraph(f"<b>Utilisation :</b> {util[:510]}..." if len(util) > 510 else f"<b>Utilisation :</b> {util}", text_style)
                    separator_top = HRFlowable(width="100%", thickness=0.5, color=black, spaceBefore=0, spaceAfter=0)

                    wrapped_util = KeepInFrame(135.73, 46, [separator_top, util_para, Spacer(1, 2)], mode='truncate')
                    util_story = [wrapped_util]
                else:
                    util_story = []

                warning_text = "<b>Avertissement !</b> Usage externe uniquement. Éviter tout contact avec les yeux. Tenir hors de portée des enfants. En cas d’apparition de rougeurs, de gonflements ou de démangeaisons pendant ou après l’utilisation, consultez un médecin. <b>A consommer de préférence avant le / Numéro de lot :</b> indiqué sur l’emballage"
                if len(warning_text) > 400:
                    warning_text = warning_text[:400] + "..."
                warning_para = Paragraph(warning_text, small_text_style)
                separator_bottom = HRFlowable(width="100%", thickness=0.5, color=black, spaceBefore=0, spaceAfter=0)
                wrapped_warning = KeepInFrame(135.73, 253, [separator_bottom, warning_para, Spacer(1, 1),separator_bottom], mode='truncate')
                warning_story = [wrapped_warning]

                vendor_text = row.get("Vendor", "")
                info_text = f"<b>Fabricant :</b> {vendor_text} EU RP : Emmanuelle Kueny - Yoomi K-Beauty, 19 rue mercière, 68100 Mulhouse, France - 03 65 67 40 62 Distributeur : ABW, 5/F, KC100, 100 Kwai Cheong Road, Kwai Chung, New territories, HongKong. <b>Fabriqué en Corée</b>"
                if len(info_text) > 400:
                    info_text = info_text[:400] + "..."
                info_para = Paragraph(info_text, small_text_style)
                wrapped_info = KeepInFrame(135.73, 25, [info_para, Spacer(1, 2)], mode='truncate')
                info_story = [wrapped_info]

                website_info = [Paragraph("www.yoomishop.fr", small_text_style)]

                # Regrouper les frames et contenus
                frames_and_stories = [
                    (Frame(3, 123,135.73, 15, showBoundary=0, leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0), title_story),
                    (Frame(3, 111, 135.73, 15, showBoundary=0, leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0), desc_story),
                    (Frame(3, 69, 135.73, 46, showBoundary=0, leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0), util_story),
                    (Frame(3, 43, 135.73, 27, showBoundary=0, leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0), warning_story),
                    (Frame(3, 20, 135.73, 25, showBoundary=0, leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0), info_story),
                    (Frame(102, -18, 135.73, 28, showBoundary=0, leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0), website_info),
                ]

                def build_all(canvas, doc):
                    for frame, story in frames_and_stories:
                        frame.addFromList(story, canvas)

                pdf.build([Spacer(0, 0)], onFirstPage=build_all)

                # Ajout des icônes avec fitz
                buffer.seek(0)
                doc = fitz.open(stream=buffer.getvalue(), filetype="pdf")
                page = doc[0]

                pao_value = row.get("custom.periode_mois", "")
                pao_icon = "pao_12m.png"
                if pd.notna(pao_value) and str(pao_value).strip() != "":
                    try:
                        pao_int = int(float(pao_value))
                        pao_icon = f"pao_{pao_int}m.png"
                    except:
                        pass

                tri_value = str(row.get("custom.texte_recyclage", "")).strip().lower().replace(" ", "_")
                tri_icon = f"{tri_value}.png" if tri_value not in ['', 'nan'] else "tri_standard.png"
                logo_icon = "logo.png"

                icon_buffer = BytesIO()
                icon_canvas = Canvas(icon_buffer, pagesize=(141.73, 141.73))
                try:
                    if os.path.exists(f"icones/{pao_icon}"):
                        icon_canvas.drawImage(f"icones/{pao_icon}", x=80, y=5, width=20, height=20)
                except:
                    pass
                try:
                    if os.path.exists(f"icones/{tri_icon}"):
                        icon_canvas.drawImage(f"icones/{tri_icon}", x=1, y=5, width=80, height=20)
                except:
                    pass
                try:
                    if os.path.exists(f"icones/{logo_icon}"):
                        icon_canvas.drawImage(f"icones/{logo_icon}", x=99, y=11, width=40, height=14)
                except:
                    pass

                icon_canvas.save()
                icon_buffer.seek(0)

                icon_doc = fitz.open(stream=icon_buffer.getvalue(), filetype="pdf")
                page.show_pdf_page(page.rect, icon_doc, 0)

                final_buffer = BytesIO()
                doc.save(final_buffer)
                final_buffer.seek(0)

                st.markdown(f"### 📰 Aperçu : {row['label']}")
                pix = page.get_pixmap(dpi=200)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                st.image(img)

                st.download_button(
                    label=f"📅 Télécharger {row['label']}.pdf",
                    data=final_buffer.getvalue(),
                    file_name=f"{row['label'].replace(' ', '_')}.pdf",
                    mime="application/pdf"
                )











with tab7:
    st.markdown("## 💸 Gestion des Soldes (manuelle par sélection)")

    if "df" in st.session_state:
        df = st.session_state["df"].copy()
        df['label_soldes'] = df['Vendor'] + ' - ' + df['Title']
        selected_soldes = st.multiselect("🛍️ Sélectionne les produits à solder", options=df['label_soldes'].tolist(), key="soldes_selection")

        # Saisie du tag à appliquer (ex : soldes30, soldes50)
        tag_to_apply = st.text_input("🏷️ Tag à appliquer (ex : soldes30)", value="soldes30")

        # Bouton pour ajouter le tag
        if st.button("✅ Ajouter le tag aux produits sélectionnés"):
            headers = {"X-Shopify-Access-Token": access_token}
            products = []
            base_url = f"https://{shop_url}/admin/api/2024-01/products.json?limit=250"
            url = base_url

            # Récupération des produits pour mapping titre → ID
            while url:
                resp = requests.get(url, headers=headers)
                resp.raise_for_status()
                products += resp.json().get("products", [])
                link = resp.headers.get("Link", "")
                if 'rel="next"' in link:
                    match = re.search(r'<([^>]+)>; rel="next"', link)
                    url = match.group(1) if match else None
                else:
                    break

            titre_to_id = {p['title']: p for p in products}

            for label in selected_soldes:
                vendor, title = label.split(" - ", 1)
                produit = titre_to_id.get(title)
                if not produit:
                    st.warning(f"❌ Produit introuvable : {title}")
                    continue

                tags_existants = produit.get("tags", "")
                nouveaux_tags = [t.strip() for t in tags_existants.split(",") if t.strip()]
                if tag_to_apply.lower() not in [t.lower() for t in nouveaux_tags]:
                    nouveaux_tags.append(tag_to_apply)

                    update_url = f"https://{shop_url}/admin/api/2024-01/products/{produit['id']}.json"
                    payload = {"product": {"id": produit["id"], "tags": ", ".join(nouveaux_tags)}}
                    update_resp = requests.put(update_url, headers=headers, json=payload)

                    if update_resp.ok:
                        st.success(f"🏷️ Tag '{tag_to_apply}' ajouté à {title}")
                    else:
                        st.error(f"❌ Erreur API sur {title} : {update_resp.text}")
                else:
                    st.info(f"ℹ️ Tag déjà présent sur {title}")

    st.markdown("## 💸 Gestion des soldes automatiques Shopify")

    if "df" not in st.session_state:
        st.warning("Charge d'abord les produits dans l’onglet 1.")
    else:
        shop_url = st.secrets["shopify"]["shop_url"]
        access_token = st.secrets["shopify"]["access_token"]
        api_version = "2024-01"

        headers = {
            "X-Shopify-Access-Token": access_token,
            "Content-Type": "application/json"
        }

        def round_up_to_0_05(value):
            return round((value * 20 + 0.9999) // 1 / 20, 2)

        def get_all_products():
            all_products = []
            url = f"https://{shop_url}/admin/api/{api_version}/products.json?limit=250"
            while url:
                resp = requests.get(url, headers=headers)
                if resp.status_code != 200:
                    st.error(f"Erreur API : {resp.status_code} - {resp.text}")
                    break
                products = resp.json().get("products", [])
                all_products.extend(products)

                link = resp.headers.get("Link", "")
                next_url = None
                if 'rel="next"' in link:
                    parts = link.split(",")
                    for part in parts:
                        if 'rel="next"' in part:
                            next_url = part.split(";")[0].strip().strip("<>").replace(" ", "")
                url = next_url
            return all_products

        def extract_discount(tags):
            for tag in tags.split(","):
                tag = tag.strip().lower()
                if tag.startswith("soldes"):
                    try:
                        return int(tag.replace("soldes", ""))
                    except:
                        return None
            return None

        def apply_discount(product, discount_percent):
            for variant in product["variants"]:
                current_price = float(variant["price"])
                compare_at = variant.get("compare_at_price")
                compare_price = float(compare_at) if compare_at else current_price

                discounted = round_up_to_0_05(compare_price * (1 - discount_percent / 100))

                needs_update = (
                    compare_at is None or
                    abs(compare_price - current_price) < 0.01 or
                    abs(current_price - discounted) > 0.01
                )

                if not needs_update:
                    continue

                variant_payload = {
                    "variant": {
                        "id": variant["id"],
                        "price": str(discounted),
                        "compare_at_price": str(compare_price)
                    }
                }

                resp = requests.put(
                    f"https://{shop_url}/admin/api/{api_version}/variants/{variant['id']}.json",
                    headers=headers,
                    json=variant_payload
                )

                if resp.ok:
                    st.success(f"✔️ {product['title']} → {compare_price}€ → {discounted}€")
                else:
                    st.error(f"❌ {product['title']} : {resp.text}")

        def revert_discount(product, soldes_tag):
            title = product["title"]
            product_id = product["id"]
            tags = product.get("tags", "")
            new_tags = [tag for tag in tags.split(",") if tag.strip().lower() != soldes_tag]

            updated = False
            for variant in product["variants"]:
                compare_at = variant.get("compare_at_price")
                if compare_at:
                    update = {
                        "variant": {
                            "id": variant["id"],
                            "price": str(compare_at),
                            "compare_at_price": None
                        }
                    }
                    resp = requests.put(
                        f"https://{shop_url}/admin/api/{api_version}/variants/{variant['id']}.json",
                        headers=headers,
                        json=update
                    )
                    if resp.ok:
                        st.success(f"♻️ {title} : retour à {compare_at}€")
                        updated = True
                    else:
                        st.error(f"❌ {title} : {resp.text}")

            if updated:
                tag_payload = {
                    "product": {
                        "id": product_id,
                        "tags": ", ".join(new_tags)
                    }
                }
                tag_resp = requests.put(
                    f"https://{shop_url}/admin/api/{api_version}/products/{product_id}.json",
                    headers=headers,
                    json=tag_payload
                )
                if tag_resp.ok:
                    st.info(f"🧹 Tag '{soldes_tag}' supprimé de {title}")
                else:
                    st.warning(f"⚠️ Tags non mis à jour pour {title}")

        if st.button("✅ Appliquer les remises selon les tags (ex: soldes30)"):
            produits = get_all_products()
            for prod in produits:
                remise = extract_discount(prod.get("tags", ""))
                if remise:
                    apply_discount(prod, remise)

        if st.button("🔁 Annuler les soldes et restaurer les prix d’origine"):
            produits = get_all_products()
            for prod in produits:
                tag_soldes = extract_discount(prod.get("tags", ""))
                if tag_soldes:
                    revert_discount(prod, f"soldes{tag_soldes}")




# --- Onglet 8 : Nouveaux produits depuis CSV + prix + poids ---
# --- Onglet 8 : Nouveaux produits depuis CSV ou TXT QUDO ---
with tab8:
    st.markdown("## ➕ Nouveaux produits — parsing + prix + poids")
    st.write(
        "Choisis la source (CSV StyleKorean **ou** TXT QUDO). "
        "Calcule les PV, crée en **brouillon**, enregistre le **coût** et le **poids**."
    )

    # --- Shopify creds
    shop_url = st.secrets["shopify"]["shop_url"]
    access_token = st.secrets["shopify"]["access_token"]
    api_version = "2024-01"
    headers = {"X-Shopify-Access-Token": access_token, "Content-Type": "application/json"}

    import math, time, requests, re

    # ---------- barcodes déjà existants ----------
    known_barcodes = set()
    if "df" in st.session_state and not st.session_state["df"].empty:
        try:
            known_barcodes = set(
                st.session_state["df"]["Variant Barcode"].astype(str).dropna().tolist()
            )
        except Exception:
            pass

    # ---------- paramètres communs ----------
    usd_to_eur_rate = st.number_input(
        "💱 Taux USD → EUR",
        value=0.92, min_value=0.5, max_value=2.0, step=0.01,
        key="usd_eur_common_tab8"
    )
    multiplier = st.number_input(
        "📈 Multiplicateur PV (ex: 2.8)",
        value=2.8, min_value=1.0, max_value=10.0, step=0.05,
        key="mult_common_tab8"
    )
    rounding_mode = st.selectbox(
        "🎯 Style d’arrondi",
        [".90 (vers le bas)", "0,10 le + proche", ".95 (vers le bas)", "arrondi sup. à 0,05", "aucun"],
        index=0,
        key="round_common_tab8"
    )

    # ---------- utilitaire de création Shopify (commun aux 2 branches) ----------
    def create_products(df_rows, default_product_type, headers, shop_url, api_version):
        progress = st.progress(0.0)
        created = 0
        total = max(1, len(df_rows))

        for _, row in df_rows.iterrows():
            try:
                title    = (row.get("Title") or "").strip() or "Sans nom"
                vendor   = (row.get("Vendor") or "").strip()
                barcode  = (row.get("Barcode") or "").strip()
                size_val = (row.get("Size") or "").strip()
                cost_eur = row.get("Cost EUR", None)
                pv_eur   = row.get("PV conseillé EUR", None)
                weight_g = row.get("Weight (g)", None)

                variant_obj = {
                    "barcode": barcode,
                    "price": f"{pv_eur:.2f}" if pd.notna(pv_eur) else "0.00",
                    "inventory_management": "shopify",
                    "inventory_policy": "deny",
                }
                if pd.notna(weight_g):
                    try:
                        variant_obj["weight"] = float(round(float(weight_g), 3))
                        variant_obj["weight_unit"] = "g"
                    except Exception:
                        pass

                product_payload = {
                    "product": {
                        "title": title,
                        "vendor": vendor,
                        "status": "draft",
                        "variants": [variant_obj],
                    }
                }
                if default_product_type and default_product_type.strip():
                    product_payload["product"]["product_type"] = default_product_type.strip()

                # 1) créer le produit
                resp = requests.post(
                    f"https://{shop_url}/admin/api/{api_version}/products.json",
                    headers=headers,
                    json=product_payload,
                )
                if resp.status_code not in (200, 201):
                    st.error(f"❌ Échec création '{title}' ({barcode}) : {resp.text}")
                    created += 1
                    progress.progress(created / total)
                    continue

                prod = resp.json().get("product", {})
                prod_id = prod.get("id")
                variant = (prod.get("variants") or [{}])[0]
                inventory_item_id = variant.get("inventory_item_id")
                st.success(f"✅ Brouillon créé : {title} — ID {prod_id}")

                # 2) metafield taille
                if size_val:
                    metafield_payload = {
                        "metafield": {
                            "namespace": "custom",
                            "key": "taille",
                            "type": "single_line_text_field",
                            "value": size_val,
                        }
                    }
                    _ = requests.post(
                        f"https://{shop_url}/admin/api/{api_version}/products/{prod_id}/metafields.json",
                        headers=headers,
                        json=metafield_payload,
                    )

                # 3) coût (EUR)
                if pd.notna(cost_eur) and inventory_item_id:
                    inv_payload = {"inventory_item": {"id": inventory_item_id, "cost": float(round(cost_eur, 2))}}
                    _ = requests.put(
                        f"https://{shop_url}/admin/api/{api_version}/inventory_items/{inventory_item_id}.json",
                        headers=headers,
                        json=inv_payload,
                    )

            except Exception as e:
                st.error(f"❌ Erreur inattendue : {e}")

            created += 1
            time.sleep(0.4)  # anti-quota
            progress.progress(created / total)

        st.success(f"🎉 Créations terminées : {created}/{total}.")

    # ---------- Sélecteur de source ----------
    source_new = st.selectbox(
        "📦 Source à créer",
        ["StyleKorean (CSV)", "QUDO (TXT)"],
        key="src_new_tab8"
    )

    # ============= BRANCHE CSV (StyleKorean) ==================
    if source_new == "StyleKorean (CSV)":
        default_vendor_csv = st.text_input(
            "🏭 Vendor par défaut (si absent dans le texte)",
            value="STYLE KOREAN",
            key="vendor_csv_tab8"
        )
        default_product_type_csv = st.text_input(
            "📦 Type de produit (optionnel)",
            value="",
            key="type_csv_tab8"
        )

        csv_new = st.file_uploader(
            "📁 CSV fournisseur (Product Name, Retail Price, Weight)",
            type=["csv"],
            key="new_csv_tab8"
        )

        if csv_new:
            try:
                df_sup = pd.read_csv(csv_new)

                # helpers locaux (CSV)
                def find_col(cands):
                    for c in df_sup.columns:
                        if str(c).strip().lower() in [x.lower() for x in cands]:
                            return c
                    return None

                col_name   = find_col(["Product Name", "name", "Nom", "Produit"])
                col_retail = find_col(["Retail Price", "Retail", "Price", "Tarif"])
                col_weight = find_col(["Weight", "Poids"])

                if not col_name:
                    st.error("❌ Colonne 'Product Name' introuvable.")
                else:
                    # — parse Product Name (ta fonction existe déjà plus haut)
                    parsed_rows = df_sup[col_name].apply(parse_product_name).tolist()
                    df_parsed = pd.DataFrame(parsed_rows)
                    df_parsed["Barcode"] = df_parsed["Barcode"].astype(str).str.extract(r'(\d{8,14})')
                    df_parsed["Vendor"] = df_parsed["Vendor"].replace("", None).fillna(default_vendor_csv)

                    # retail → cost USD
                    def extract_usd_from_retail(retail_raw: str):
                        if pd.isna(retail_raw):
                            return None
                        s = str(retail_raw)
                        lines = [x.strip() for x in re.split(r'[\r\n]+', s) if x.strip()]
                        nums = []
                        for ln in lines:
                            m = re.search(r'(\d[\d,]*\.?\d*)', ln)
                            if m:
                                try:
                                    nums.append(float(m.group(1).replace(",", "")))
                                except Exception:
                                    pass
                        if len(nums) >= 2:
                            return nums[1]
                        return nums[0] if nums else None

                    if col_retail:
                        df_parsed["Cost USD"] = df_sup[col_retail].apply(extract_usd_from_retail)
                    else:
                        df_parsed["Cost USD"] = None

                    # conversions + PV conseillés (ta fonction price_rounding existe déjà au-dessus)
                    df_parsed["Cost EUR"]         = df_parsed["Cost USD"].apply(lambda x: round(x * usd_to_eur_rate, 2) if pd.notna(x) else None)
                    df_parsed["PV brut EUR"]      = df_parsed["Cost EUR"].apply(lambda x: round(x * multiplier, 2) if pd.notna(x) else None)
                    df_parsed["PV conseillé EUR"] = df_parsed["PV brut EUR"].apply(lambda x: price_rounding(x, rounding_mode) if pd.notna(x) else None)

                    # poids (g)
                    if col_weight:
                        df_parsed["Weight (g)"] = df_sup[col_weight].apply(parse_weight_to_grams)
                    else:
                        df_parsed["Weight (g)"] = None

                    # nouveaux produits
                    df_new = df_parsed[df_parsed["Barcode"].notna() & ~df_parsed["Barcode"].isin(known_barcodes)].copy()
                    st.dataframe(
                        df_new[["Vendor","Title","Size","Barcode","Weight (g)","Cost USD","Cost EUR","PV conseillé EUR"]],
                        use_container_width=True
                    )

                    labels_csv = (
                        df_new["Vendor"].astype(str) + " — " +
                        df_new["Title"].astype(str) + " — " +
                        df_new["Size"].astype(str) + " — " +
                        df_new["Barcode"].astype(str)
                    ).tolist()
                    to_create_csv = st.multiselect(
                        "Sélectionne les produits à créer",
                        options=labels_csv,
                        key="sel_csv_tab8"
                    )

                    if st.button("🧪 Créer en brouillon (CSV)", key="btn_create_csv_tab8"):
                        sel = df_new[
                            (df_new["Vendor"].astype(str) + " — " +
                             df_new["Title"].astype(str) + " — " +
                             df_new["Size"].astype(str) + " — " +
                             df_new["Barcode"].astype(str)).isin(to_create_csv)
                        ]
                        if sel.empty:
                            st.warning("Aucune sélection.")
                        else:
                            create_products(sel, default_product_type_csv, headers, shop_url, api_version)

            except Exception as e:
                st.error(f"Erreur lecture/traitement CSV : {e}")

    # ============= BRANCHE TXT (QUDO) ==================
    else:
        # NOTE: parse_qudo_text_to_df & parse_qudo_name doivent être définies plus haut (helpers)
        default_vendor_txt = st.text_input(
            "🏭 Vendor par défaut (si absent dans le nom)",
            value="",
            key="vendor_txt_tab8"
        )
        default_product_type_txt = st.text_input(
            "📦 Type de produit (optionnel)",
            value="",
            key="type_txt_tab8"
        )
        default_weight_g = st.number_input(
            "⚖️ Poids (g) par défaut",
            value=0.0, min_value=0.0, step=1.0,
            key="weight_txt_tab8"
        )

        txt_new = st.file_uploader(
            "📄 Fichier texte QUDO",
            type=["txt"],
            key="new_txt_tab8"
        )

        if txt_new:
            try:
                # 1) Parse TXT QUDO -> DataFrame (contient déjà "Unit Price EUR")
                content = txt_new.read().decode("utf-8", errors="ignore")
                df_txt = parse_qudo_text_to_df(content, include_samples=False)

                # 2) Découpe Vendor / Title / Size à partir du Product Name
                parsed = df_txt["Product Name"].apply(lambda x: parse_qudo_name(x, default_vendor_txt)).apply(pd.Series)
                df_parsed = pd.concat([parsed, df_txt[["Barcode","Unit Price EUR"]]], axis=1)
                df_parsed["Barcode"] = df_parsed["Barcode"].astype(str).str.extract(r'(\d{8,14})')
                df_parsed = df_parsed[df_parsed["Barcode"].notna()]

                # 3) Coût = prix QUDO (déjà en EUR) ; poids par défaut
                df_parsed["Cost EUR"]   = df_parsed["Unit Price EUR"].astype(float)
                df_parsed["Weight (g)"] = default_weight_g

                # 4) Calcul PV conseillé (à partir de Cost EUR) avec tes paramètres globaux multiplier/rounding_mode
                df_parsed["PV brut EUR"]      = df_parsed["Cost EUR"].apply(lambda x: round(x * multiplier, 2) if pd.notna(x) else None)
                df_parsed["PV conseillé EUR"] = df_parsed["PV brut EUR"].apply(lambda x: price_rounding(x, rounding_mode) if pd.notna(x) else None)

                # 5) Retirer ce qui existe déjà dans Shopify (barcodes connus)
                df_new = df_parsed[~df_parsed["Barcode"].isin(known_barcodes)].copy()

                st.dataframe(
                    df_new[["Vendor","Title","Size","Barcode","Weight (g)","Cost EUR","PV conseillé EUR"]],
                    use_container_width=True
                )

                # 6) Sélection des produits à créer
                labels_txt = (
                    df_new["Vendor"].astype(str) + " — " +
                    df_new["Title"].astype(str) + " — " +
                    df_new["Size"].astype(str) + " — " +
                    df_new["Barcode"].astype(str)
                ).tolist()
                to_create_txt = st.multiselect(
                    "Sélectionne les produits à créer",
                    options=labels_txt,
                    key="sel_txt_tab8"
                )

                # 7) Création en brouillon sur Shopify
                if st.button("🧪 Créer en brouillon (TXT QUDO)", key="btn_create_txt_tab8"):
                    sel = df_new[
                        (df_new["Vendor"].astype(str) + " — " +
                         df_new["Title"].astype(str) + " — " +
                         df_new["Size"].astype(str) + " — " +
                         df_new["Barcode"].astype(str)).isin(to_create_txt)
                    ]
                    if sel.empty:
                        st.warning("Aucune sélection.")
                    else:
                        create_products(sel, default_product_type_txt, headers, shop_url, api_version)

            except Exception as e:
                st.error(f"Erreur lecture/traitement TXT : {e}")

